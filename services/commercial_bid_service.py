"""商务标 / 资格标响应：基于 tender_detail 生成章节并落库，支持确认态保护与分册导出。"""

from __future__ import annotations

import hashlib
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from config import OUTPUT_DIR
from db.migrations import BID_SCOPE_TECHNICAL, BID_SCOPE_TECHNICAL_COMMERCIAL
from db.models import CommercialSection, Project
from services.tender_detail_service import get_tender_detail
from services.word_styling import apply_professional_styles

SECTION_NOTICE = "notice"
SECTION_COMMERCE_REQUIREMENT = "commerce_requirement"
SECTION_QUALIFICATION = "qualification"
SECTION_COMMERCE_SCORE = "commerce_score"

STATUS_DRAFT = "draft"
STATUS_CONFIRMED = "confirmed"


def _md_escape_title(title: str) -> str:
    return (title or "").strip() or "未命名"


def _stable_match_key(*parts: Any) -> str:
    raw = "|".join(str(p or "").strip() for p in parts)
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:16]


def get_bid_scope(project: Project) -> str:
    scope = (getattr(project, "bid_scope", None) or BID_SCOPE_TECHNICAL).strip()
    if scope not in (BID_SCOPE_TECHNICAL, BID_SCOPE_TECHNICAL_COMMERCIAL):
        return BID_SCOPE_TECHNICAL
    return scope


def set_bid_scope(project: Project, enabled: bool) -> str:
    project.bid_scope = (
        BID_SCOPE_TECHNICAL_COMMERCIAL if enabled else BID_SCOPE_TECHNICAL
    )
    return project.bid_scope


def _build_draft_sections(project: Project) -> list[dict[str, Any]]:
    detail = get_tender_detail(project)
    notice = detail.get("notice") or {}
    commerce_req = (detail.get("commerce_requirements") or "").strip()
    qual_items = detail.get("qualification_items") or []
    commerce_scores = detail.get("commerce_scores") or []

    sections: list[dict[str, Any]] = []
    sort_order = 0

    notice_md = "\n".join([
        "## 投标人须知响应摘要",
        "",
        f"- 工程名称：{project.name or notice.get('project_name') or '待填'}",
        f"- 合同模式：{notice.get('contract_mode') or '待填'}",
        f"- 暗标：{'是' if notice.get('blind_bid') is True else '否' if notice.get('blind_bid') is False else '未明确'}",
        "",
    ])
    sections.append({
        "section_key": SECTION_NOTICE,
        "match_key": _stable_match_key(SECTION_NOTICE, "summary"),
        "title": "投标人须知响应摘要",
        "content_markdown": notice_md,
        "sort_order": sort_order,
    })
    sort_order += 1

    if commerce_req:
        req_md = "\n".join(["## 商务要求逐条响应", "", commerce_req, ""])
    else:
        req_md = "\n".join([
            "## 商务要求逐条响应",
            "",
            "（招标文件未提取到商务要求正文，请人工补充）",
            "",
        ])
    sections.append({
        "section_key": SECTION_COMMERCE_REQUIREMENT,
        "match_key": _stable_match_key(SECTION_COMMERCE_REQUIREMENT, "body"),
        "title": "商务要求逐条响应",
        "content_markdown": req_md,
        "sort_order": sort_order,
    })
    sort_order += 1

    if qual_items:
        for i, item in enumerate(qual_items, start=1):
            label = _md_escape_title(item.get("item_label") or "资格审查")
            desc = (item.get("description") or "").strip()
            seq = item.get("seq") or i
            md = "\n".join([
                f"### {seq}. {label}",
                "",
                f"**招标要求：** {desc or '（无）'}",
                "",
                "**响应说明：** 完全响应，详见证明材料。",
                "",
            ])
            sections.append({
                "section_key": SECTION_QUALIFICATION,
                "match_key": _stable_match_key(
                    SECTION_QUALIFICATION, seq, label, desc[:80],
                ),
                "title": f"资格审查：{label}",
                "content_markdown": md,
                "sort_order": sort_order,
            })
            sort_order += 1
    else:
        sections.append({
            "section_key": SECTION_QUALIFICATION,
            "match_key": _stable_match_key(SECTION_QUALIFICATION, "empty"),
            "title": "资格审查响应",
            "content_markdown": "（未提取到资格审查条目）\n",
            "sort_order": sort_order,
        })
        sort_order += 1

    if commerce_scores:
        for i, score in enumerate(commerce_scores, start=1):
            title = _md_escape_title(score.get("title"))
            criteria = (score.get("criteria") or "").strip()
            value = score.get("score_value")
            score_hint = f"（{value} 分）" if value not in (None, "") else ""
            lines = [f"### {i}. {title}{score_hint}", ""]
            if criteria:
                lines.extend([f"**评分标准：** {criteria}", ""])
            lines.extend([
                f"**响应要点：** 针对「{title}」逐条响应评分标准，附证明材料索引；"
                "不得偏离招标文件实质性要求。",
                "",
            ])
            sections.append({
                "section_key": SECTION_COMMERCE_SCORE,
                "match_key": _stable_match_key(
                    SECTION_COMMERCE_SCORE, title, criteria[:80], value,
                ),
                "title": f"商务评分：{title}",
                "content_markdown": "\n".join(lines),
                "sort_order": sort_order,
            })
            sort_order += 1
    else:
        sections.append({
            "section_key": SECTION_COMMERCE_SCORE,
            "match_key": _stable_match_key(SECTION_COMMERCE_SCORE, "empty"),
            "title": "商务评分项响应",
            "content_markdown": "（未提取到商务评分项）\n",
            "sort_order": sort_order,
        })

    return sections


def list_commercial_sections(db: Session, project_id: str) -> list[CommercialSection]:
    return (
        db.query(CommercialSection)
        .filter(CommercialSection.project_id == project_id)
        .order_by(CommercialSection.sort_order, CommercialSection.id)
        .all()
    )


def section_to_dict(row: CommercialSection) -> dict[str, Any]:
    return {
        "id": row.id,
        "project_id": row.project_id,
        "section_key": row.section_key,
        "match_key": row.match_key,
        "title": row.title,
        "content_markdown": row.content_markdown,
        "status": row.status,
        "sort_order": row.sort_order,
        "updated_at": row.updated_at.isoformat() if row.updated_at else None,
    }


def persist_commercial_draft(
    db: Session,
    project: Project,
    *,
    preserve_confirmed: bool = True,
) -> list[CommercialSection]:
    """按 tender_detail 生成/刷新 CommercialSection；confirmed 默认保留。"""
    drafts = _build_draft_sections(project)
    existing = {
        (row.section_key, row.match_key): row
        for row in list_commercial_sections(db, project.id)
    }
    keep_ids: set[str] = set()
    result: list[CommercialSection] = []

    for item in drafts:
        key = (item["section_key"], item["match_key"])
        old = existing.get(key)
        if old and preserve_confirmed and old.status == STATUS_CONFIRMED:
            old.sort_order = item["sort_order"]
            keep_ids.add(old.id)
            result.append(old)
            continue
        if old:
            old.title = item["title"]
            old.content_markdown = item["content_markdown"]
            old.sort_order = item["sort_order"]
            if not preserve_confirmed or old.status != STATUS_CONFIRMED:
                old.status = STATUS_DRAFT
            keep_ids.add(old.id)
            result.append(old)
        else:
            row = CommercialSection(
                project_id=project.id,
                section_key=item["section_key"],
                match_key=item["match_key"],
                title=item["title"],
                content_markdown=item["content_markdown"],
                status=STATUS_DRAFT,
                sort_order=item["sort_order"],
            )
            db.add(row)
            result.append(row)

    for old in existing.values():
        if old.id in keep_ids:
            continue
        if preserve_confirmed and old.status == STATUS_CONFIRMED:
            keep_ids.add(old.id)
            result.append(old)
            continue
        db.delete(old)

    db.flush()
    return sorted(result, key=lambda r: (r.sort_order, r.id))


def update_commercial_section(
    db: Session,
    project_id: str,
    section_id: str,
    *,
    title: str | None = None,
    content_markdown: str | None = None,
    status: str | None = None,
) -> CommercialSection:
    row = (
        db.query(CommercialSection)
        .filter(
            CommercialSection.project_id == project_id,
            CommercialSection.id == section_id,
        )
        .first()
    )
    if not row:
        raise ValueError("商务标章节不存在")
    if title is not None:
        row.title = title.strip() or row.title
    if content_markdown is not None:
        row.content_markdown = content_markdown
    if status is not None:
        if status not in (STATUS_DRAFT, STATUS_CONFIRMED):
            raise ValueError("无效的章节状态")
        row.status = status
    db.flush()
    return row


def commercial_status(db: Session, project: Project) -> dict[str, Any]:
    rows = list_commercial_sections(db, project.id)
    by_key: dict[str, dict[str, int]] = {}
    for row in rows:
        bucket = by_key.setdefault(row.section_key, {"total": 0, "confirmed": 0})
        bucket["total"] += 1
        if row.status == STATUS_CONFIRMED:
            bucket["confirmed"] += 1
    return {
        "bid_scope": get_bid_scope(project),
        "enabled": get_bid_scope(project) == BID_SCOPE_TECHNICAL_COMMERCIAL,
        "section_count": len(rows),
        "confirmed_count": sum(1 for r in rows if r.status == STATUS_CONFIRMED),
        "by_section_key": by_key,
        "sections": [section_to_dict(r) for r in rows],
    }


def assemble_commercial_markdown(db: Session, project: Project) -> str:
    rows = list_commercial_sections(db, project.id)
    if not rows:
        return build_commercial_draft(project)["markdown"]
    lines = [
        f"# {project.name or '工程'} 商务与资格响应文件（草稿）",
        "",
    ]
    for row in rows:
        lines.append(row.content_markdown.rstrip())
        lines.append("")
    lines.extend([
        "## 声明",
        "",
        "本响应文件为系统根据招标文件解析结果自动生成的草稿，正式递交前须由商务人员核对并盖章。",
        "",
    ])
    return "\n".join(lines)


def build_commercial_draft(project: Project) -> dict[str, Any]:
    """兼容旧接口：返回整份 Markdown 摘要（不落库）。"""
    detail = get_tender_detail(project)
    sections = _build_draft_sections(project)
    lines = [
        f"# {project.name or '工程'} 商务与资格响应文件（草稿）",
        "",
    ]
    for item in sections:
        lines.append(item["content_markdown"].rstrip())
        lines.append("")
    lines.extend([
        "## 声明",
        "",
        "本响应文件为系统根据招标文件解析结果自动生成的草稿，正式递交前须由商务人员核对并盖章。",
        "",
    ])
    return {
        "markdown": "\n".join(lines),
        "qualification_count": len(detail.get("qualification_items") or []),
        "commerce_score_count": len(detail.get("commerce_scores") or []),
        "has_commerce_requirements": bool((detail.get("commerce_requirements") or "").strip()),
        "section_count": len(sections),
    }


def export_commercial_docx(project: Project, db: Session | None = None) -> Path:
    if db is not None:
        md = assemble_commercial_markdown(db, project)
    else:
        md = build_commercial_draft(project)["markdown"]
    out_dir = Path(OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r'[\\/:*?"<>|]', "_", project.name or "商务资格响应")
    date_tag = datetime.now().strftime("%Y%m%d")
    out_path = out_dir / f"{safe_name}_商务资格_{date_tag}.docx"

    doc = Document()
    for line in md.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("|"):
            continue
        else:
            doc.add_paragraph(stripped)

    detail = get_tender_detail(project)
    qual_items = detail.get("qualification_items") or []
    if qual_items:
        doc.add_heading("资格审查响应表（附表）", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["序号", "审查类别", "招标要求", "响应说明"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for i, item in enumerate(qual_items, start=1):
            row = table.add_row().cells
            row[0].text = str(item.get("seq") or i)
            row[1].text = str(item.get("item_label") or "")
            row[2].text = str(item.get("description") or "")
            row[3].text = "完全响应，详见证明材料"

    apply_professional_styles(doc)
    doc.save(str(out_path))
    return out_path
