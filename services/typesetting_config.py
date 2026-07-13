"""自定义排版参数：标题/正文字体、编号、对齐等，导出 Word 时生效。"""

from __future__ import annotations

from copy import deepcopy
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from db.models import Project
from services.generation_config import get_generation_config
from services.numbering_service import TWIPS_PER_LEVEL, _next_numbering_id

LEVEL_KEYS = ("h1", "h2", "h3", "h4", "h5", "body")
HEADING_LEVEL_KEYS = ("h1", "h2", "h3", "h4", "h5")
MAX_TYPESETTING_HEADING_LEVELS = 5

FONT_OPTIONS = ("黑体", "宋体", "楷体", "仿宋", "微软雅黑")
FONT_SIZE_OPTIONS = (
    "二号", "小二", "三号", "小三", "四号", "小四", "五号", "小五",
)
ALIGN_OPTIONS = ("left", "center", "right", "justify")
NUMBER_FORMAT_OPTIONS = (
    "cn_dun",
    "cn_paren_dun",
    "decimal_dun",
    "decimal_1_2",
    "decimal_1_2_3",
    "decimal_full",
    "decimal",
    "none",
)

FONT_SIZE_PT: dict[str, float] = {
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
}

ALIGN_LABELS = {
    "left": "左对齐",
    "center": "居中",
    "right": "右对齐",
    "justify": "两端对齐",
}

NUMBER_FORMAT_LABELS = {
    "cn_dun": "一、二、三、",
    "cn_paren_dun": "（一）、（二）、",
    "decimal_dun": "1、2、3、",
    "decimal_1_2": "1.1 1.2",
    "decimal_1_2_3": "1.1.1",
    "decimal_full": "1.1.1.1",
    "decimal": "1 2 3",
    "none": "—",
}

# (numFmt, lvlText) — lvlText 中 %N 对应 OOXML 多级编号占位
_NUMBER_FORMAT_DEFS: dict[str, tuple[str, str] | None] = {
    "cn_dun": ("chineseCounting", "%1、"),
    "cn_paren_dun": ("chineseCounting", "（%2）、"),
    "decimal_dun": ("decimal", "%3、"),
    "decimal_1_2": ("decimal", "%1.%2"),
    "decimal_1_2_3": ("decimal", "%1.%2.%3"),
    "decimal_full": ("decimal", "%1.%2.%3.%4"),
    "decimal": ("decimal", "%1"),
    "none": None,
}

_STYLE_TO_LEVEL = {
    "Title": "h1",
    "Heading 1": "h1",
    "Heading 2": "h2",
    "Heading 3": "h3",
    "Heading 4": "h4",
    "Heading 5": "h5",
}

_DEFAULT_LEVEL = {
    "number_format": "decimal",
    "font": "宋体",
    "font_size": "小四",
    "align": "left",
    "bold": False,
    "color": "#000000",
    "first_line_indent": 0,
}

DEFAULT_TYPESETTING: dict[str, dict[str, Any]] = {
    "h1": {
        "number_format": "cn_dun",
        "font": "黑体",
        "font_size": "小三",
        "align": "center",
        "bold": True,
        "color": "#000000",
        "first_line_indent": 0,
    },
    "h2": {
        "number_format": "cn_paren_dun",
        "font": "宋体",
        "font_size": "小四",
        "align": "left",
        "bold": True,
        "color": "#000000",
        "first_line_indent": 0,
    },
    "h3": {
        "number_format": "decimal_dun",
        "font": "宋体",
        "font_size": "小四",
        "align": "left",
        "bold": True,
        "color": "#000000",
        "first_line_indent": 0,
    },
    "h4": {
        "number_format": "decimal_1_2",
        "font": "宋体",
        "font_size": "小四",
        "align": "left",
        "bold": False,
        "color": "#000000",
        "first_line_indent": 0,
    },
    "h5": {
        "number_format": "decimal_full",
        "font": "宋体",
        "font_size": "小四",
        "align": "left",
        "bold": False,
        "color": "#000000",
        "first_line_indent": 0,
    },
    "body": {
        "number_format": "none",
        "font": "宋体",
        "font_size": "小四",
        "align": "left",
        "bold": False,
        "color": "#000000",
        "first_line_indent": 2,
    },
}

TABLE_HEADER_BG = "191970"


def list_typesetting_options() -> dict[str, Any]:
    return {
        "levels": [
            {"key": k, "label": _level_label(k)}
            for k in LEVEL_KEYS
        ],
        "fonts": list(FONT_OPTIONS),
        "font_sizes": list(FONT_SIZE_OPTIONS),
        "aligns": [{"value": k, "label": ALIGN_LABELS[k]} for k in ALIGN_OPTIONS],
        "number_formats": [
            {"value": k, "label": NUMBER_FORMAT_LABELS[k]}
            for k in NUMBER_FORMAT_OPTIONS
        ],
        "defaults": default_typesetting(),
    }


def _level_label(key: str) -> str:
    labels = {
        "h1": "一级标题",
        "h2": "二级标题",
        "h3": "三级标题",
        "h4": "四级标题",
        "h5": "五级标题",
        "body": "正文",
    }
    return labels.get(key, key)


def default_typesetting() -> dict[str, dict[str, Any]]:
    return deepcopy(DEFAULT_TYPESETTING)


def _normalize_level(raw: dict[str, Any] | None, *, level_key: str) -> dict[str, Any]:
    base = deepcopy(DEFAULT_TYPESETTING.get(level_key, _DEFAULT_LEVEL))
    if isinstance(raw, dict):
        base.update({k: v for k, v in raw.items() if k in base})
    if base["font"] not in FONT_OPTIONS:
        base["font"] = "宋体"
    if base["font_size"] not in FONT_SIZE_PT:
        base["font_size"] = "小四"
    if base["align"] not in ALIGN_OPTIONS:
        base["align"] = "left"
    if base["number_format"] not in NUMBER_FORMAT_OPTIONS:
        base["number_format"] = "none" if level_key == "body" else DEFAULT_TYPESETTING.get(level_key, _DEFAULT_LEVEL)["number_format"]
    base["bold"] = bool(base.get("bold"))
    color = str(base.get("color") or "#000000").strip()
    if not color.startswith("#") or len(color) != 7:
        color = "#000000"
    base["color"] = color
    try:
        base["first_line_indent"] = max(0, min(8, int(base.get("first_line_indent") or 0)))
    except (TypeError, ValueError):
        base["first_line_indent"] = 0
    return base


def normalize_typesetting(raw: dict[str, Any] | None) -> dict[str, dict[str, Any]]:
    result: dict[str, dict[str, Any]] = {}
    src = raw if isinstance(raw, dict) else {}
    for key in LEVEL_KEYS:
        result[key] = _normalize_level(src.get(key), level_key=key)
    return result


def get_typesetting(project: Project) -> dict[str, dict[str, Any]]:
    cfg = get_generation_config(project)
    return normalize_typesetting(cfg.get("typesetting"))


def _hex_to_rgb(color: str) -> RGBColor:
    color = color.lstrip("#")
    r = int(color[0:2], 16)
    g = int(color[2:4], 16)
    b = int(color[4:6], 16)
    return RGBColor(r, g, b)


def _align_enum(align: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def _set_run_font(
    run,
    font_name: str,
    size_pt: float,
    *,
    bold: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = _hex_to_rgb(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _apply_level_style(paragraph: Paragraph, level_cfg: dict[str, Any]) -> None:
    font = level_cfg["font"]
    size_pt = FONT_SIZE_PT[level_cfg["font_size"]]
    bold = level_cfg["bold"]
    color = level_cfg["color"]
    paragraph.alignment = _align_enum(level_cfg["align"])
    pf = paragraph.paragraph_format
    indent_chars = level_cfg["first_line_indent"]
    if indent_chars > 0:
        pf.first_line_indent = Pt(size_pt * indent_chars)
    else:
        pf.first_line_indent = Pt(0)
    if level_cfg.get("_is_body"):
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in paragraph.runs:
        _set_run_font(run, font, size_pt, bold=bold, color=color)


def _make_custom_lvl(ilvl: int, num_fmt: str, lvl_text: str) -> OxmlElement:
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(ilvl))

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt_el = OxmlElement("w:numFmt")
    num_fmt_el.set(qn("w:val"), num_fmt)
    lvl.append(num_fmt_el)

    lvl_text_el = OxmlElement("w:lvlText")
    lvl_text_el.set(qn("w:val"), lvl_text)
    lvl.append(lvl_text_el)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    if ilvl > 0:
        lvl_restart = OxmlElement("w:lvlRestart")
        lvl_restart.set(qn("w:val"), str(ilvl - 1))
        lvl.append(lvl_restart)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(TWIPS_PER_LEVEL * ilvl))
    ind.set(qn("w:hanging"), str(TWIPS_PER_LEVEL))
    p_pr.append(ind)
    lvl.append(p_pr)
    return lvl


def _register_custom_numbering(doc: Document, typesetting: dict[str, dict[str, Any]]) -> int | None:
    """按各级编号格式注册多级列表，返回 numId；若各级均为 none 则返回 None。"""
    level_defs: list[tuple[int, str, str]] = []
    for ilvl, key in enumerate(HEADING_LEVEL_KEYS):
        fmt_id = typesetting[key]["number_format"]
        fmt = _NUMBER_FORMAT_DEFS.get(fmt_id)
        if fmt is None:
            continue
        level_defs.append((ilvl, fmt[0], fmt[1]))

    if not level_defs:
        return None

    numbering_elm = doc.part.numbering_part.element
    abstract_num_id = _next_numbering_id(numbering_elm, "w:abstractNum", qn("w:abstractNumId"))
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)

    for ilvl, num_fmt, lvl_text in level_defs:
        abstract_num.append(_make_custom_lvl(ilvl, num_fmt, lvl_text))

    numbering_elm.append(abstract_num)

    num_id = _next_numbering_id(numbering_elm, "w:num", qn("w:numId"))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id_el = OxmlElement("w:abstractNumId")
    abstract_num_id_el.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_num_id_el)
    numbering_elm.append(num)
    return num_id


class TypesettingNumbering:
    """按自定义排版参数挂载 Word 多级编号。"""

    def __init__(self, doc: Document, typesetting: dict[str, dict[str, Any]]):
        self._num_id = _register_custom_numbering(doc, typesetting)
        self._active_ilvls = set()
        for ilvl, key in enumerate(HEADING_LEVEL_KEYS):
            if typesetting[key]["number_format"] != "none":
                self._active_ilvls.add(ilvl)

    @property
    def enabled(self) -> bool:
        return self._num_id is not None

    def apply(self, paragraph: Paragraph, outline_level: int) -> None:
        if self._num_id is None:
            return
        ilvl = min(max(outline_level, 1), MAX_TYPESETTING_HEADING_LEVELS) - 1
        if ilvl not in self._active_ilvls:
            return
        p_pr = paragraph._p.get_or_add_pPr()
        if p_pr.find(qn("w:numPr")) is not None:
            return
        num_pr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_pr.append(ilvl_el)
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(self._num_id))
        num_pr.append(num_id_el)
        p_pr.append(num_pr)


def apply_document_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(72)


def apply_typesetting_styles(doc: Document, typesetting: dict[str, dict[str, Any]] | None) -> None:
    """按自定义排版参数统一段落版式。"""
    ts = normalize_typesetting(typesetting)
    apply_document_margins(doc)

    for para in doc.paragraphs:
        style_name = (para.style.name if para.style else "") or ""
        text = (para.text or "").strip()
        if not text and not para.runs:
            continue

        level_key = _STYLE_TO_LEVEL.get(style_name)
        if not level_key:
            if style_name in ("Normal", "List Bullet", "List Number") or not style_name.startswith("Heading"):
                level_key = "body"
            else:
                continue

        level_cfg = dict(ts[level_key])
        level_cfg["_is_body"] = level_key == "body"
        _apply_level_style(para, level_cfg)

    body_cfg = ts["body"]
    body_size = FONT_SIZE_PT[body_cfg["font_size"]]
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if row_idx == 0:
                            _set_run_font(
                                run, "黑体", 10, bold=True, color="#FFFFFF",
                            )
                        else:
                            _set_run_font(
                                run,
                                body_cfg["font"],
                                body_size,
                                bold=body_cfg["bold"],
                                color=body_cfg["color"],
                            )
                if row_idx == 0:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), TABLE_HEADER_BG)
                    tc_pr.append(shd)
