import json
import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from chart.chart_service import (
    CHART_PATTERN,
    next_caption,
    parse_chart_match,
    render_flow,
    render_gantt,
    render_org,
    render_timeline,
    render_warning_image,
)
from config import OUTPUT_DIR
from db.models import Project, TechOutline
from services.blind_bid_service import anonymize_cover_meta, blind_header_text, is_blind_bid
from services.numbering_service import HeadingNumbering, resolve_heading_numbering_preset
from services.word_styling import (
    add_blind_bid_header,
    apply_professional_styles,
    append_toc_and_body_sections,
    create_cover_document,
    enable_auto_update_fields,
)

logger = logging.getLogger(__name__)

STYLE_MAP = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
HIGHLIGHT_PATTERN = re.compile(r"\*\*\[([^\]]+)\]\s*([^*]+)\*\*")
HEADING_LINE_PATTERN = re.compile(r"^(#+)\s+(.+)$")
_TABLE_ROW_RE = re.compile(r"^\|.*\|$")
_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")
_UL_ITEM_RE = re.compile(r"^[-*+]\s+(.*)$")
_OL_ITEM_RE = re.compile(r"^\d+[.、]\s+(.*)$")


def _render_cover(project: Project) -> Document:
    meta = {
        "project_date": datetime.now().strftime("%Y年%m月%d日"),
        "voltage_level": project.voltage_level or "",
        "capacity": project.capacity or "",
        "location": project.location or "",
    }
    if is_blind_bid(project):
        meta = anonymize_cover_meta(meta)
    return create_cover_document(project.name or "技术方案", meta)


def _add_highlight_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    _add_highlight_runs(para, text)


def _insert_smart_table(doc: Document, data: list[dict]) -> bool:
    """插入 SMART_DATA 对照表。data 为空时不插入，返回 False。"""
    if not data:
        return False
    table = doc.add_table(rows=2, cols=len(data))
    table.style = "Table Grid"
    for i, item in enumerate(data):
        table.rows[0].cells[i].text = item.get("title", "")
        table.rows[1].cells[i].text = item.get("desc", "")
    return True


def _insert_picture_centered(doc: Document, img: Path, width: Pt) -> None:
    doc.add_picture(str(img), width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _insert_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def _next_caption_number(counters: dict, kind: str) -> int:
    counters[kind] = counters.get(kind, 0) + 1
    return counters[kind]


def _split_lines_at_first_list(lines: list[str]) -> tuple[list[str], list[str]] | None:
    """首行非列表、后续出现列表时，拆成导语行 + 列表行。"""
    for i, line in enumerate(lines):
        if i == 0:
            continue
        stripped = line.strip()
        if stripped and _detect_list_style(stripped):
            return lines[:i], lines[i:]
    return None


def _insert_chart_image(
    doc: Document, img: Path, width: Pt, temp_files: list[Path], counters: dict, chart_type: str
) -> None:
    _insert_picture_centered(doc, img, width)
    temp_files.append(img)
    _insert_caption(doc, next_caption(counters, chart_type))


def _insert_chart(
    doc: Document, match: re.Match, temp_files: list[Path], duration: int, counters: dict
) -> None:
    chart_type, raw_json = parse_chart_match(match)
    try:
        data = json.loads(raw_json)
        if chart_type == "GANTT_DATA":
            img = render_gantt(data, duration)
            _insert_chart_image(doc, img, Pt(450), temp_files, counters, chart_type)
        elif chart_type == "TIMELINE_DATA":
            img = render_timeline(data)
            _insert_chart_image(doc, img, Pt(400), temp_files, counters, chart_type)
        elif chart_type == "FLOW_DATA":
            img = render_flow(data) or render_warning_image("工艺流程图")
            _insert_chart_image(doc, img, Pt(400), temp_files, counters, chart_type)
        elif chart_type == "ORG_DATA":
            img = render_org(data) or render_warning_image("组织架构图")
            _insert_chart_image(doc, img, Pt(350), temp_files, counters, chart_type)
        elif chart_type == "SMART_DATA":
            if _insert_smart_table(doc, data):
                _insert_caption(doc, next_caption(counters, chart_type))

    except Exception as exc:
        logger.warning("图表解析失败 %s: %s", chart_type, exc)
        img = render_warning_image(chart_type)
        _insert_picture_centered(doc, img, Pt(400))
        temp_files.append(img)


from services.markdown_utils import parse_table_cells

_parse_table_cells = parse_table_cells


def _is_markdown_table(lines: list[str]) -> bool:
    """判断开头两行是否是 Markdown 表格（表头行 + 分隔行）。"""
    if len(lines) < 2:
        return False
    return bool(_TABLE_ROW_RE.match(lines[0]) and _TABLE_SEP_RE.match(lines[1]))


def _insert_markdown_table(doc: Document, lines: list[str]) -> int:
    """把开头连续的 Markdown 表格行转换成真正的 Word 表格，返回消耗的行数。"""
    header = parse_table_cells(lines[0])
    col_count = max(len(header), 1)

    data_rows: list[list[str]] = []
    consumed = 2  # 表头行 + 分隔行
    for line in lines[2:]:
        if not _TABLE_ROW_RE.match(line):
            break
        data_rows.append(parse_table_cells(line))
        consumed += 1

    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = "Table Grid"

    for i, cell in enumerate(table.rows[0].cells):
        cell.text = header[i] if i < len(header) else ""
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for ri, row in enumerate(data_rows, start=1):
        for ci, cell in enumerate(table.rows[ri].cells):
            cell.text = row[ci] if ci < len(row) else ""

    return consumed


def _detect_list_style(line: str) -> str | None:
    if _UL_ITEM_RE.match(line):
        return "List Bullet"
    if _OL_ITEM_RE.match(line):
        return "List Number"
    return None


def _add_highlight_runs(para, text: str) -> None:
    last = 0
    for m in HIGHLIGHT_PATTERN.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        run = para.add_run(f"{m.group(1)} {m.group(2)}")
        run.bold = True
        run.font.highlight_color = 7  # yellow
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def _insert_markdown_list(doc: Document, lines: list[str], style: str) -> int:
    """把开头连续、同一种列表标记的 Markdown 列表行转换成 Word 列表段落，返回消耗的行数。"""
    pattern = _UL_ITEM_RE if style == "List Bullet" else _OL_ITEM_RE
    consumed = 0
    for line in lines:
        m = pattern.match(line.strip())
        if not m:
            break
        item_text = m.group(1).strip()
        para = doc.add_paragraph(style=style)
        if "**[" in item_text:
            _add_highlight_runs(para, item_text)
        else:
            para.add_run(item_text)
        consumed += 1
    return consumed


def _write_text_block(doc: Document, text: str, chapter_level: int) -> None:
    for para_text in text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        lines = para_text.split("\n")
        first_line = lines[0].strip()
        heading_match = HEADING_LINE_PATTERN.match(first_line)
        if heading_match:
            hash_count = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            doc.add_heading(title, level=min(chapter_level + hash_count, 4))
            rest_lines = [l for l in lines[1:] if l.strip()]
            if rest_lines:
                _write_paragraph_or_table(doc, rest_lines)
            continue
        _write_paragraph_or_table(doc, lines)


def _write_paragraph_or_table(doc: Document, lines: list[str]) -> None:
    """把一段文本行写入文档：依次识别 Markdown 表格 / 列表（可能后面还跟着普通段落），
    否则按普通段落/高亮段落写入。"""
    stripped = [l.strip() for l in lines]
    if _is_markdown_table(stripped):
        consumed = _insert_markdown_table(doc, stripped)
        remaining = [l for l in lines[consumed:] if l.strip()]
        if remaining:
            _write_paragraph_or_table(doc, remaining)
        return

    list_style = _detect_list_style(stripped[0]) if stripped and stripped[0] else None
    if list_style:
        consumed = _insert_markdown_list(doc, lines, list_style)
        remaining = [l for l in lines[consumed:] if l.strip()]
        if remaining:
            _write_paragraph_or_table(doc, remaining)
        return

    intro_split = _split_lines_at_first_list(lines)
    if intro_split:
        intro_lines, list_lines = intro_split
        intro_text = "\n".join(intro_lines).strip()
        if intro_text:
            if "**[" in intro_text:
                _add_highlight_paragraph(doc, intro_text)
            else:
                doc.add_paragraph(intro_text)
        list_style = _detect_list_style(list_lines[0].strip())
        if list_style:
            consumed = _insert_markdown_list(doc, list_lines, list_style)
            remaining = [l for l in list_lines[consumed:] if l.strip()]
            if remaining:
                _write_paragraph_or_table(doc, remaining)
            return

    text = "\n".join(lines).strip()
    if "**[" in text:
        _add_highlight_paragraph(doc, text)
    else:
        doc.add_paragraph(text)


def _write_content_in_order(
    doc: Document, content: str, chapter_level: int, temp_files: list[Path], duration: int, counters: dict
) -> None:
    last_end = 0
    for match in CHART_PATTERN.finditer(content):
        _write_text_block(doc, content[last_end:match.start()], chapter_level)
        _insert_chart(doc, match, temp_files, duration, counters)
        last_end = match.end()
    _write_text_block(doc, content[last_end:], chapter_level)


def _should_emit_heading(chapter: TechOutline) -> bool:
    """叶子、一级章、或有自身正文的节点输出标题；跳过无内容的中间容器标题。"""
    if chapter.is_leaf == 1:
        return True
    if (chapter.generated_content or "").strip():
        return True
    if chapter.level == 1:
        return True
    return False


def _compute_heading_emit_levels(chapters: list[TechOutline]) -> dict[str, int]:
    """按实际输出的标题重算编号层级，跳过无内容中间节点时不留空号。"""
    by_id = {ch.id: ch for ch in chapters}
    levels: dict[str, int] = {}

    def ancestor_emit_level(chapter: TechOutline) -> int:
        parent_id = chapter.parent_id
        while parent_id:
            if parent_id in levels:
                return levels[parent_id]
            parent = by_id.get(parent_id)
            parent_id = parent.parent_id if parent else None
        return 0

    for ch in chapters:
        if not _should_emit_heading(ch):
            continue
        levels[ch.id] = ancestor_emit_level(ch) + 1
    return levels


def assemble_document(
    project: Project,
    chapters: list[TechOutline],
    *,
    mark_yellow: bool = False,
) -> Path:
    out_dir = Path(OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r'[\\/:*?"<>|]', "_", project.name or "技术方案")
    date_tag = datetime.now().strftime("%Y%m%d")
    out_path = out_dir / f"{safe_name}_{date_tag}.docx"

    doc = _render_cover(project)
    append_toc_and_body_sections(doc)
    if is_blind_bid(project):
        for section in doc.sections:
            add_blind_bid_header(section, blind_header_text())
    heading_numbering = HeadingNumbering(doc, resolve_heading_numbering_preset(project))
    emit_levels = _compute_heading_emit_levels(chapters)
    temp_files: list[Path] = []
    duration = project.duration_days or 90
    counters: dict = {}
    seen_level1 = False

    try:
        for ch in chapters:
            if not _should_emit_heading(ch):
                continue
            if ch.level == 1:
                if seen_level1:
                    doc.add_page_break()
                seen_level1 = True
            style = STYLE_MAP.get(ch.level, "Heading 2")
            heading_title = ch.title
            if mark_yellow and ch.is_leaf == 1 and ch.review_status == "yellow":
                heading_title = f"{ch.title}【待优化】"
            emit_level = emit_levels.get(ch.id, ch.level)
            try:
                heading_para = doc.add_heading(heading_title, level=min(emit_level, 4))
            except Exception:
                heading_para = doc.add_paragraph(heading_title, style=style)
            heading_numbering.apply(heading_para, emit_level)

            content = ch.generated_content or ""
            if content.strip():
                _write_content_in_order(doc, content, emit_level, temp_files, duration, counters)

        apply_professional_styles(doc)
        enable_auto_update_fields(doc)
        doc.save(str(out_path))
    finally:
        for path in temp_files:
            try:
                path.unlink(missing_ok=True)
            except OSError as exc:
                logger.warning("临时图表文件清理失败 %s: %s", path, exc)

    return out_path
