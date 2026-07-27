"""整本合规终审 + 跨章节一致性检查（借鉴 tender-writer-v4 阶段 5）。"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from db.models import Project, TechOutline, TechRequirement
from services.project_meta import get_meta, set_meta
from services.outline_order import sort_outline_tree_dfs
from services.qa_rules import (
    SUBSTANTIAL_KEYWORDS,
    check_font_safety,
    check_template_residues,
    extract_coverage_candidates,
    keyword_covered_in_headings,
    mandatory_element_covered,
    normalize_for_match,
    split_keywords,
    split_mandatory_elements,
)

_MONEY_RE = re.compile(r"(\d+(?:\.\d+)?)\s*万元|(\d{4,})\s*元")
from services.duration_text import DURATION_DAYS_RE as _DAYS_RE
_D_PLUS_RE = re.compile(r"D\s*\+\s*(\d+)")
_PERSON_RE = re.compile(r"(\d+)\s*(人|位)")
_TEAM_CONTEXT = ("项目组", "团队", "项目经理", "技术骨干", "核心人员", "专家", "工程师", "成员", "人员配置", "项目部")
_RESUME_KW = ("简历", "师资", "团队成员", "人员配置", "项目人员", "骨干")
_ORG_KW = ("组织架构", "组织结构", "项目组织", "人员架构", "团队架构", "团队组成")
_NAME_RE = re.compile(
    r"([一-龥]{2,4})"
    r"(?:"
    r"[,，、]\s*(?:担任|为|出任|是|负责)?\s*"
    r"|(?:担任|为|出任|是|负责)\s*)"
    r"([\u4e00-\u9fa5]{2,8})"
)


def read_docx_text(docx_path: Path) -> tuple[str, dict]:
    from docx import Document

    document = Document(str(docx_path))
    paragraph_texts: list[str] = []
    paragraph_styles: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraph_texts.append(text)
            paragraph_styles.append(getattr(para.style, "name", ""))
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraph_texts.append(" | ".join(cells))
    format_info = {
        "paragraph_count": len(document.paragraphs),
        "non_empty_paragraphs": len(paragraph_texts),
        "heading_count": sum(1 for name in paragraph_styles if name.startswith("Heading")),
        "cover_preview": paragraph_texts[:12],
    }
    return "\n".join(paragraph_texts), format_info


def _chapters_combined_text(chapters: list[TechOutline]) -> str:
    parts: list[str] = []
    for ch in sorted(chapters, key=lambda c: c.sort_order):
        parts.append(ch.title or "")
        if ch.generated_content:
            parts.append(ch.generated_content)
    return "\n".join(parts)


def _parse_money_yuan(text: str) -> list[float]:
    out: list[float] = []
    for m in _MONEY_RE.finditer(text):
        if m.group(1):
            try:
                out.append(float(m.group(1)) * 10000)
            except ValueError:
                pass
        elif m.group(2):
            try:
                out.append(float(m.group(2)))
            except ValueError:
                pass
    return out


def check_cross_consistency(
    project: Project,
    docx_text: str,
    meta: dict[str, Any],
) -> list[dict[str, str]]:
    items: list[dict[str, str]] = []
    duration = project.duration_days
    budget = meta.get("budget_yuan")
    if budget is not None:
        try:
            budget = float(budget)
        except (TypeError, ValueError):
            budget = None

    d_plus = [int(m.group(1)) for m in _D_PLUS_RE.finditer(docx_text)]
    if duration and d_plus:
        hard = [d for d in sorted(set(d_plus)) if d > duration * 2]
        soft = [d for d in sorted(set(d_plus)) if duration < d <= duration * 2]
        for d in hard:
            items.append({
                "level": "fail",
                "message": f"D+{d} 明显超出总工期 {duration} 天（超 {d - duration} 天）",
            })
        for d in soft:
            items.append({
                "level": "warn",
                "message": f"D+{d} 超出总工期 {duration} 天，若属评审配合节点请人工复核",
            })
        if not hard and not soft:
            items.append({"level": "pass", "message": f"D+N 节点均在 {duration} 天工期内"})

    if duration:
        for num, _ in _DAYS_RE.findall(docx_text):
            if int(num) > duration * 2:
                items.append({
                    "level": "warn",
                    "message": f"正文工期数字 {num} 天与全局 {duration} 天偏差过大",
                })

    people: list[int] = []
    for m in _PERSON_RE.finditer(docx_text):
        ctx = docx_text[max(0, m.start() - 20): m.start()]
        if any(w in ctx for w in _TEAM_CONTEXT):
            people.append(int(m.group(1)))

    if people and duration and budget and budget > 0:
        max_people = max(people)
        estimated = max_people * 0.5 * duration * 1500
        threshold = budget * 1.5
        level = "fail" if estimated > threshold else "pass"
        items.append({
            "level": level,
            "message": (
                f"团队成本估算 {estimated:,.0f} 元（{max_people} 人×50%×{duration} 天×1500 元/人·天），"
                f"合同预算 {budget:,.0f} 元"
            ),
        })

    if budget and budget > 0:
        doc_moneys = _parse_money_yuan(docx_text)
        hard = [m for m in doc_moneys if m > budget * 10]
        if hard:
            items.append({
                "level": "fail",
                "message": f"正文金额 {hard[:3]} 元大幅超出预算（10 倍以上）",
            })

    resume_names = _extract_section_names(docx_text, _RESUME_KW)
    org_names = _extract_section_names(docx_text, _ORG_KW)
    if resume_names and org_names:
        only_resume = sorted(resume_names - org_names)
        if only_resume:
            items.append({
                "level": "warn",
                "message": f"简历章节有、组织架构未出现的人名：{', '.join(only_resume[:8])}",
            })

    return items


def _extract_section_names(
    text: str, section_keywords: tuple[str, ...], *, max_lines: int = 60,
) -> set[str]:
    """扫描简历/组织架构类小节内的人名；用行数窗口兜底关闭，不依赖 `#`（正文/docx 文本均无此符号）。"""
    lines = text.splitlines()
    in_section = False
    section_line_count = 0
    names: set[str] = set()
    for line in lines:
        stripped = line.strip()
        if any(kw in line for kw in section_keywords) and len(line) < 40:
            in_section = True
            section_line_count = 0
            continue
        if in_section:
            section_line_count += 1
            # 空行 + 短行大概率是新小节标题，或超过窗口长度，判定小节结束
            if (not stripped and section_line_count > 3) or section_line_count > max_lines:
                in_section = False
                continue
            for m in _NAME_RE.finditer(line):
                names.add(m.group(1))
    return names


def check_scoring_coverage(
    docx_text: str,
    requirements: list[TechRequirement],
) -> list[dict[str, Any]]:
    normalized_doc = normalize_for_match(docx_text)
    results: list[dict[str, Any]] = []
    for req in requirements:
        if req.status == "ignored":
            continue
        candidates = extract_coverage_candidates(req.requirement_title, req.keyword)
        exact = normalize_for_match(req.requirement_title) in normalized_doc
        matched = [c for c in candidates if normalize_for_match(c) in normalized_doc]
        if exact or len(matched) >= 2:
            status = "covered"
        elif len(matched) == 1:
            status = "partial"
        else:
            status = "missing"
        results.append({
            "id": req.id,
            "title": req.requirement_title,
            "score": req.score_value,
            "status": status,
            "matched": matched,
        })
    return results


def check_substantial_response(docx_text: str, requirements: list[TechRequirement]) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for req in requirements:
        if req.is_risk_item != 1 and "★" not in (req.requirement_title or "") and "▲" not in (req.requirement_title or ""):
            risk = req.risk_hint or ""
            if "★" not in risk and "▲" not in risk:
                continue
        candidates = extract_coverage_candidates(req.requirement_title, req.keyword)
        responded = False
        evidence = ""
        for candidate in candidates:
            idx = 0
            while True:
                pos = docx_text.find(candidate, idx)
                if pos == -1:
                    break
                window = docx_text[max(0, pos - 80): pos + len(candidate) + 120]
                phrase = next((p for p in SUBSTANTIAL_KEYWORDS if p in window), "")
                if phrase:
                    responded = True
                    evidence = f"{candidate} 附近出现「{phrase}」"
                    break
                idx = pos + len(candidate)
            if responded:
                break
        results.append({
            "title": req.requirement_title,
            "responded": responded,
            "evidence": evidence,
        })
    return results


def check_title_keywords_from_outline(
    chapters: list[TechOutline],
    requirements: list[TechRequirement],
) -> list[dict[str, Any]]:
    req_map = {r.id: r for r in requirements}
    issues: list[dict[str, Any]] = []
    for ch in chapters:
        if ch.is_leaf != 1:
            continue
        import json

        req_ids = []
        try:
            req_ids = json.loads(ch.requirement_ids or "[]")
        except json.JSONDecodeError:
            pass
        keywords: list[str] = []
        for rid in req_ids:
            req = req_map.get(rid)
            if req:
                keywords.extend(split_keywords(req.keyword))
        keywords = list(dict.fromkeys(keywords))
        if not keywords:
            continue
        titles = [ch.title or ""]
        hit = [kw for kw in keywords if keyword_covered_in_headings(titles, kw)]
        if not hit:
            issues.append({"chapter": ch.title, "missing_keywords": keywords})
    return issues


def check_mandatory_elements_doc(
    docx_text: str,
    requirements: list[TechRequirement],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for req in requirements:
        if req.status == "ignored":
            continue
        elements = split_mandatory_elements(req.mandatory_elements)
        if not elements:
            continue
        missing = [e for e in elements if not mandatory_element_covered(docx_text, e)]
        if missing:
            issues.append({
                "title": req.requirement_title,
                "missing": missing,
            })
    return issues


def check_chapter_length_balance(
    chapters: list[TechOutline],
    *,
    threshold: float = 0.4,
) -> list[dict[str, str]]:
    """检测单章篇幅是否占全文过高比例。"""
    leaves = [
        ch for ch in chapters
        if ch.is_leaf == 1 and (ch.generated_content or "").strip()
    ]
    if len(leaves) < 2:
        return []

    lengths = {ch.id: len(ch.generated_content or "") for ch in leaves}
    total = sum(lengths.values())
    if total <= 0:
        return []

    issues: list[dict[str, str]] = []
    for ch in leaves:
        ratio = lengths[ch.id] / total
        if ratio > threshold:
            issues.append(
                {
                    "level": "warn",
                    "message": (
                        f"章节「{ch.title}」约占全文 {ratio:.0%}"
                        f"（{lengths[ch.id]} 字），篇幅分布失衡"
                    ),
                }
            )
    return issues


def _render_markdown_report(
    sections: dict[str, Any],
    passed: bool,
    *,
    category_summary: dict[str, Any] | None = None,
) -> str:
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines = [
        "# 合规终审报告",
        "",
        f"- 生成时间：{now}",
        f"- 总体结论：**{'通过' if passed else '未通过'}**",
        "",
    ]
    if category_summary:
        lines.append("## 检查分类汇总")
        lines.append("")
        by_cat = category_summary.get("by_category") or {}
        if not by_cat:
            lines.append("（无问题）")
        else:
            for cat, payload in by_cat.items():
                label = payload.get("label") or cat
                count = payload.get("count", 0)
                lines.append(f"- **{label}**：{count} 项")
        lines.append("")
    for title, payload in sections.items():
        lines.append(f"## {title}")
        lines.append("")
        if isinstance(payload, list):
            if not payload:
                lines.append("（无问题）")
            for item in payload:
                if isinstance(item, dict):
                    level = item.get("level", item.get("status", item.get("severity", "")))
                    msg = item.get("message") or item.get("title") or str(item)
                    prefix = {
                        "fail": "✗",
                        "block": "✗",
                        "warn": "⚠",
                        "pass": "✓",
                        "info": "·",
                        "missing": "✗",
                        "partial": "△",
                    }.get(level, "·")
                    lines.append(f"- {prefix} {msg}")
                else:
                    lines.append(f"- {item}")
        else:
            lines.append(str(payload))
        lines.append("")
    return "\n".join(lines)


def run_compliance(
    db: Session,
    project: Project,
    docx_path: Path | None,
    chapters: list[TechOutline],
) -> dict[str, Any]:
    requirements = (
        db.query(TechRequirement)
        .filter(TechRequirement.project_id == project.id)
        .order_by(TechRequirement.is_risk_item.desc())
        .all()
    )
    meta = get_meta(project)
    if docx_path and docx_path.exists():
        docx_text, format_info = read_docx_text(docx_path)
    else:
        docx_text = _chapters_combined_text(chapters)
        format_info = {}

    from services.check_catalog import category_label
    from services.check_registry import ProjectCheckContext, run_project_checks, summarize_findings
    from services.tender_detail_service import filter_qualification_items, get_tender_detail

    detail = get_tender_detail(project)
    qualification_items = filter_qualification_items(
        detail.get("qualification_items") or [],
        "废标项",
    )
    # 若无显式「废标项」tab 数据，回退全部资格/符合性条目供对照
    if not qualification_items:
        qualification_items = detail.get("qualification_items") or []

    findings = run_project_checks(
        ProjectCheckContext(
            project=project,
            chapters=chapters,
            requirements=requirements,
            docx_text=docx_text,
            meta=meta,
            docx_path=docx_path,
            format_info=format_info,
            qualification_items=qualification_items,
        )
    )
    summary = summarize_findings(findings)
    fail_count = summary["block_count"]
    # info 不计入 warning；与历史 warn 语义对齐
    warn_count = summary["warn_count"]

    # 保留原 sections 结构供前端/测试兼容；按分类重组 + 保留明细
    cross = check_cross_consistency(project, docx_text, meta)
    coverage = check_scoring_coverage(docx_text, requirements)
    substantial = check_substantial_response(docx_text, requirements)
    residues = check_template_residues(docx_text)
    title_kw = check_title_keywords_from_outline(chapters, requirements)
    mandatory = check_mandatory_elements_doc(docx_text, requirements)
    length_balance = check_chapter_length_balance(chapters)
    font_issues = check_font_safety(docx_path) if docx_path and docx_path.exists() else []
    disqual_findings = [f for f in findings if f.category == "disqualification_risk"]

    sections = {
        "一、跨章节一致性": cross,
        "二、评分项覆盖度": [
            f"{'✗' if c['status'] == 'missing' else '△' if c['status'] == 'partial' else '✓'} "
            f"{c['title']}（{c['status']}）"
            for c in coverage
        ],
        "三、刚性/实质性响应": [
            f"{'✓' if s['responded'] else '✗'} {s['title']}" + (f" — {s['evidence']}" if s.get("evidence") else "")
            for s in substantial
        ],
        "四、模板残留": residues or ["（无）"],
        "五、大纲标题关键词": [
            f"✗ {i['chapter']} 缺 {', '.join(i['missing_keywords'])}" for i in title_kw
        ] or ["（无）"],
        "六、必备要素": [
            f"✗ {i['title']} 缺 {', '.join(i['missing'])}" for i in mandatory
        ] or ["（无）"],
        "七、篇幅分布": [
            f"⚠ {i['message']}" for i in length_balance
        ] or ["（无）"],
        "八、字体规范": font_issues or ["（无）"],
        "九、废标条款对照": [
            {
                "level": f.severity if f.severity != "block" else "fail",
                "message": f.message,
            }
            for f in disqual_findings
        ] or ["（无废标条款或无需对照）"],
        "十、格式摘要": [
            f"非空段落 {format_info.get('non_empty_paragraphs', '—')}，标题 {format_info.get('heading_count', '—')}",
        ],
    }

    # 分类视图（按 catalog 标签）
    category_sections: dict[str, list[dict[str, Any]]] = {}
    for f in findings:
        label = category_label(f.category)
        category_sections.setdefault(label, []).append(
            {
                "level": f.severity,
                "message": f.message,
                "check_id": f.check_id,
                "evidence": f.evidence,
            }
        )

    passed = fail_count == 0
    markdown = _render_markdown_report(
        sections,
        passed,
        category_summary=summary,
    )
    report = {
        "passed": passed,
        "failure_count": fail_count,
        "warning_count": warn_count,
        "sections": sections,
        "coverage": coverage,
        "findings": [f.to_dict() for f in findings],
        "category_summary": summary,
        "category_sections": category_sections,
        "markdown": markdown,
        "checked_at": datetime.now(timezone.utc).isoformat(),
    }
    set_meta(project, last_compliance_report=report)
    db.commit()
    return report


def get_last_compliance_report(project: Project) -> dict[str, Any] | None:
    return get_meta(project).get("last_compliance_report")


def check_compliance_now(db: Session, project: Project) -> dict[str, Any]:
    """在不导出 docx 的情况下运行合规检查（用于生成过程中随时自检）。"""
    chapters = sort_outline_tree_dfs(
        db.query(TechOutline).filter(TechOutline.project_id == project.id).all()
    )
    return run_compliance(db, project, docx_path=None, chapters=chapters)


def is_compliance_report_stale(
    db: Session,
    project: Project,
    report: dict[str, Any] | None,
) -> bool:
    """报告生成后是否有章节被重新生成过，用于前端提示"过期"。"""
    if not report or not report.get("checked_at"):
        return False
    try:
        checked_at = datetime.fromisoformat(report["checked_at"])
    except (TypeError, ValueError):
        return True
    if checked_at.tzinfo is None:
        checked_at = checked_at.replace(tzinfo=timezone.utc)
    latest = (
        db.query(TechOutline)
        .filter(TechOutline.project_id == project.id, TechOutline.generated_at.isnot(None))
        .order_by(TechOutline.generated_at.desc())
        .first()
    )
    if not latest or not latest.generated_at:
        return False
    gen_at = latest.generated_at
    if gen_at.tzinfo is None:
        gen_at = gen_at.replace(tzinfo=timezone.utc)
    return gen_at > checked_at
