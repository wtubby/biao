"""电力 EPC 技术标 Word 专业版式（借鉴 AI-Bid-System professional 工程深蓝模板）。"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 深蓝工程风（professional template）
HEADING_COLOR = RGBColor(25, 25, 112)
BODY_FONT = "宋体"
HEADING_FONT = "黑体"
SUBHEADING_FONT = "楷体"
TABLE_HEADER_BG = "191970"  # midnight blue hex without #


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def apply_professional_styles(doc: Document) -> None:
    """对已有文档统一应用电力工程技术标版式。"""
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(72)

    for para in doc.paragraphs:
        style_name = (para.style.name if para.style else "") or ""
        text = (para.text or "").strip()
        if not text and not para.runs:
            continue

        if style_name.startswith("Heading") or style_name == "Title":
            level = 0
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
            size = {0: 22, 1: 16, 2: 14, 3: 12}.get(level, 12)
            font = HEADING_FONT if level <= 1 else SUBHEADING_FONT
            for run in para.runs:
                _set_run_font(run, font, size, bold=True, color=HEADING_COLOR)
            if level == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        for run in para.runs:
            _set_run_font(run, BODY_FONT, 12)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.first_line_indent = Pt(24)

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if row_idx == 0:
                            _set_run_font(run, HEADING_FONT, 10, bold=True, color=RGBColor(255, 255, 255))
                        else:
                            _set_run_font(run, BODY_FONT, 10)
                if row_idx == 0:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), TABLE_HEADER_BG)
                    tc_pr.append(shd)


def _append_field_run(paragraph, instr: str) -> None:
    """向段落插入 Word 域代码（目录、页码等）。"""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_el)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def insert_toc_field(paragraph) -> None:
    _append_field_run(paragraph, r'TOC \o "1-3" \h \z \u')


def add_page_number_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_field_run(paragraph, "PAGE")


def add_blind_bid_header(section, text: str = "技术方案（暗标）") -> None:
    """暗标统一页眉：匿名、无公司名。"""
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        paragraph.runs[0].text = text
        _set_run_font(paragraph.runs[0], BODY_FONT, 9, color=RGBColor(89, 89, 89))
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, BODY_FONT, 9, color=RGBColor(89, 89, 89))


def set_section_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def enable_auto_update_fields(doc: Document) -> None:
    """标记文档在打开时自动重算所有域（目录、页码等），
    确保 LibreOffice headless 转 PDF 时目录不为空。"""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def append_toc_and_body_sections(doc: Document) -> None:
    """封面后插入目录节与正文节（正文节启用页码）。"""
    doc.add_page_break()
    toc_heading = doc.add_heading("目  录", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_para = doc.add_paragraph()
    insert_toc_field(toc_para)
    hint = doc.add_paragraph("（目录将在打开文档时自动更新）")
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hint.runs:
        _set_run_font(run, BODY_FONT, 9, color=RGBColor(128, 128, 128))

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_page_number_start(body_section, 1)
    add_page_number_footer(body_section)


def create_cover_document(project_name: str, meta: dict) -> Document:
    """创建带电力工程封面样式的空文档。"""
    doc = Document()
    is_blind = bool((meta or {}).get("blind_bid"))
    cover_title = "技术方案（暗标）" if is_blind else (project_name or "技术方案")
    title = doc.add_heading(cover_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_lines = []
    if meta.get("project_date"):
        info_lines.append(f"编制日期：{meta['project_date']}")
    if not is_blind:
        if meta.get("voltage_level"):
            info_lines.append(f"电压等级：{meta['voltage_level']}")
        if meta.get("capacity"):
            info_lines.append(f"工程规模：{meta['capacity']}")
        if meta.get("location"):
            info_lines.append(f"建设地点：{meta['location']}")
    else:
        # 暗标封面仅保留不可识别的技术参数档位
        if meta.get("voltage_level"):
            info_lines.append(f"电压等级：{meta['voltage_level']}")
        info_lines.append("（暗标文件，不含投标人标识）")
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    apply_professional_styles(doc)
    return doc
