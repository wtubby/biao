import logging
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber

from services.env_check import check_ghostscript

logger = logging.getLogger(__name__)

SCANNED_PDF_THRESHOLD = 100
CAMELOT_ACCURACY_THRESHOLD = 70
HEADER_FOOTER_MARGIN = 50
MIN_PARAGRAPH_LEN = 20
OCR_RENDER_SCALE = 2.0


@dataclass
class ParsedItem:
    text: str
    page: int
    kind: str  # "table" | "paragraph"


@dataclass
class ParsedContent:
    items: list[ParsedItem] = field(default_factory=list)
    page_count: int = 0
    is_scanned: bool = False
    error: str | None = None

    @property
    def tables(self) -> list[str]:
        return [i.text for i in self.items if i.kind == "table"]

    @property
    def paragraphs(self) -> list[str]:
        return [i.text for i in self.items if i.kind == "paragraph"]


def _dedupe_items(items: list[ParsedItem]) -> list[ParsedItem]:
    seen: set[str] = set()
    result: list[ParsedItem] = []
    for item in items:
        normalized = item.text.strip()
        if normalized and normalized not in seen:
            seen.add(normalized)
            result.append(item)
    return result


def _obj_center_in_bboxes(obj: dict, bboxes: list[tuple[float, float, float, float]], *, pad: float = 1.0) -> bool:
    """判断字符/对象中心是否落在任一表格 bbox 内。"""
    x0 = float(obj.get("x0") or 0)
    top = float(obj.get("top") or 0)
    x1 = float(obj.get("x1") or x0)
    bottom = float(obj.get("bottom") or top)
    cx = (x0 + x1) / 2
    cy = (top + bottom) / 2
    for bx0, btop, bx1, bbottom in bboxes:
        if (bx0 - pad) <= cx <= (bx1 + pad) and (btop - pad) <= cy <= (bbottom + pad):
            return True
    return False


def _extract_text_excluding_tables(page, table_bboxes: list[tuple[float, float, float, float]]) -> str:
    """提取正文，排除已识别表格区域内的字符，避免表格文字再进段落。"""
    if not table_bboxes:
        return page.extract_text() or ""
    filtered = page.filter(lambda obj: not _obj_center_in_bboxes(obj, table_bboxes))
    return filtered.extract_text() or ""


def _parse_pdf_camelot(file_path: Path) -> list[ParsedItem]:
    """Camelot Lattice 模式提取有线框表格。"""
    items: list[ParsedItem] = []
    try:
        import camelot

        pages = camelot.read_pdf(str(file_path), flavor="lattice", pages="all")
        for table in pages:
            if table.accuracy > CAMELOT_ACCURACY_THRESHOLD:
                text = table.df.to_string(index=False)
                page = int(table.page) if table.page else 1
                items.append(ParsedItem(text=text, page=page, kind="table"))
    except Exception as exc:
        logger.warning("Camelot 解析失败: %s", exc)
    return items


def _parse_pdf_pdfplumber(
    file_path: Path,
    *,
    skip_table_pages: set[int] | None = None,
) -> list[ParsedItem]:
    """pdfplumber 提取无线框表格和正文段落。

    表格区域会从段落提取中排除，避免同一表格内容以 table + paragraph 双重进入结果。
    """
    items: list[ParsedItem] = []
    skip_pages = skip_table_pages or set()

    with pdfplumber.open(file_path) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            cropped = page.within_bbox(
                (0, HEADER_FOOTER_MARGIN, page.width, page.height - HEADER_FOOTER_MARGIN)
            )

            found_tables = []
            try:
                found_tables = list(cropped.find_tables() or [])
            except Exception as exc:
                logger.debug("pdfplumber find_tables 失败 page=%s: %s", page_idx, exc)

            table_bboxes = [tuple(t.bbox) for t in found_tables if getattr(t, "bbox", None)]

            if page_idx not in skip_pages:
                for table in found_tables:
                    try:
                        data = table.extract()
                    except Exception:
                        data = None
                    if not data:
                        continue
                    rows = ["\t".join(cell or "" for cell in row) for row in data if row]
                    if rows:
                        items.append(
                            ParsedItem(text="\n".join(rows), page=page_idx, kind="table")
                        )

            # 即使本页表格已由 Camelot 提取，仍排除表格区域，避免段落复读表格文字
            text = _extract_text_excluding_tables(cropped, table_bboxes)
            for para in re.split(r"\n\s*\n", text):
                para = para.strip()
                if len(para) >= MIN_PARAGRAPH_LEN:
                    items.append(ParsedItem(text=para, page=page_idx, kind="paragraph"))

    return items


def _health_check_pdf(file_path: Path) -> tuple[int, bool]:
    """PyMuPDF 健康检查：判断是否为扫描件。"""
    doc = fitz.open(file_path)
    total_chars = sum(len(page.get_text()) for page in doc)
    page_count = len(doc)
    doc.close()
    is_scanned = total_chars < SCANNED_PDF_THRESHOLD
    return page_count, is_scanned


def _parse_pdf_ocr(file_path: Path) -> list[ParsedItem]:
    """扫描件 PDF 逐页 OCR（依赖 rapidocr-onnxruntime，可选安装）。"""
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        logger.warning("rapidocr-onnxruntime 未安装，无法对扫描件执行 OCR")
        return []

    ocr = RapidOCR()
    items: list[ParsedItem] = []
    doc = fitz.open(file_path)
    try:
        for page_idx, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(OCR_RENDER_SCALE, OCR_RENDER_SCALE))
            tmp_path: Path | None = None
            try:
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(pix.tobytes("png"))
                    tmp_path = Path(tmp.name)
                result, _ = ocr(str(tmp_path))
            finally:
                if tmp_path:
                    tmp_path.unlink(missing_ok=True)

            if not result:
                continue
            lines = [str(item[1]).strip() for item in result if len(item) > 1 and str(item[1]).strip()]
            page_text = "\n".join(lines)
            if len(page_text) >= MIN_PARAGRAPH_LEN:
                items.append(ParsedItem(text=page_text, page=page_idx, kind="paragraph"))
    finally:
        doc.close()
    return items


def parse_pdf(file_path: Path) -> ParsedContent:
    page_count, is_scanned = _health_check_pdf(file_path)
    if is_scanned:
        ocr_items = _parse_pdf_ocr(file_path)
        if ocr_items:
            logger.info("扫描件 PDF 已通过 OCR 提取 %d 段文本（共 %d 页）", len(ocr_items), page_count)
            return ParsedContent(
                items=_dedupe_items(ocr_items),
                page_count=page_count,
                is_scanned=True,
            )
        return ParsedContent(
            page_count=page_count,
            is_scanned=True,
            error=(
                "检测到扫描件 PDF（可提取文字不足 100 字）。"
                "请安装 OCR 依赖后重试（pip install rapidocr-onnxruntime），"
                "或提供带文字层的 PDF 文件。"
            ),
        )

    all_items: list[ParsedItem] = []
    camelot_table_pages: set[int] = set()

    if check_ghostscript():
        camelot_items = _parse_pdf_camelot(file_path)
        all_items.extend(camelot_items)
        camelot_table_pages = {item.page for item in camelot_items}
        if not camelot_items:
            logger.info("Camelot 未提取到表格，pdfplumber 将继续提取表格")
    else:
        logger.info("Ghostscript 不可用，跳过 Camelot，使用 pdfplumber 全量解析")

    all_items.extend(
        _parse_pdf_pdfplumber(file_path, skip_table_pages=camelot_table_pages)
    )

    return ParsedContent(
        items=_dedupe_items(all_items),
        page_count=page_count,
    )


def parse_docx(file_path: Path) -> ParsedContent:
    from docx import Document

    doc = Document(file_path)
    items: list[ParsedItem] = []

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        if rows:
            items.append(ParsedItem(text="\n".join(rows), page=1, kind="table"))

    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) >= MIN_PARAGRAPH_LEN:
            items.append(ParsedItem(text=text, page=1, kind="paragraph"))

    return ParsedContent(
        items=_dedupe_items(items),
        page_count=0,
    )


def parse_document(file_path: Path) -> ParsedContent:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix == ".docx":
        return parse_docx(file_path)
    if suffix == ".doc":
        return ParsedContent(error="暂不支持 .doc 格式，请先转换为 .docx 后再上传")
    return ParsedContent(error=f"不支持的文件格式: {suffix}")
