"""商务标落库与 regenerate 保护。"""

import pytest
from docx import Document

from db.database import SessionLocal, init_db
from db.migrations import BID_SCOPE_TECHNICAL_COMMERCIAL
from db.models import Project
from services.commercial_bid_service import (
    STATUS_CONFIRMED,
    STATUS_DRAFT,
    _append_commercial_markdown_to_doc,
    get_bid_scope,
    list_commercial_sections,
    persist_commercial_draft,
    set_bid_scope,
    update_commercial_section,
)
from services.tender_detail_service import set_tender_detail


def _seed_project(db):
    project = Project(name="商务标测试工程", bid_scope="technical")
    db.add(project)
    db.flush()
    set_tender_detail(
        project,
        {
            "notice": {"project_name": "商务标测试工程", "contract_mode": "EPC", "blind_bid": False},
            "commerce_requirements": "投标保证金 10 万元。",
            "qualification_items": [
                {"seq": 1, "item_label": "资格性审查", "description": "具备承装资质"},
                {"seq": 2, "item_label": "符合性审查", "description": "响应文件齐全"},
            ],
            "commerce_scores": [
                {"title": "业绩", "criteria": "近三年同类业绩", "score_value": 10},
            ],
        },
    )
    db.commit()
    db.refresh(project)
    return project


def test_persist_commercial_sections_counts():
    init_db()
    db = SessionLocal()
    try:
        project = _seed_project(db)
        set_bid_scope(project, True)
        rows = persist_commercial_draft(db, project)
        db.commit()
        assert get_bid_scope(project) == BID_SCOPE_TECHNICAL_COMMERCIAL
        # notice + commerce_requirement + 2 qual + 1 score
        assert len(rows) == 5
        keys = [r.section_key for r in rows]
        assert keys.count("qualification") == 2
        assert keys.count("commerce_score") == 1
    finally:
        db.close()


def test_regenerate_preserves_confirmed():
    init_db()
    db = SessionLocal()
    try:
        project = _seed_project(db)
        set_bid_scope(project, True)
        rows = persist_commercial_draft(db, project)
        db.commit()
        target = next(r for r in rows if r.section_key == "commerce_requirement")
        update_commercial_section(
            db,
            project.id,
            target.id,
            content_markdown="## 人工确认内容\n\n已核对。\n",
            status=STATUS_CONFIRMED,
        )
        db.commit()

        # 改 tender_detail 后 regenerate
        set_tender_detail(
            project,
            {
                "notice": {"project_name": "商务标测试工程", "contract_mode": "EPC"},
                "commerce_requirements": "投标保证金改为 20 万元。",
                "qualification_items": [
                    {"seq": 1, "item_label": "资格性审查", "description": "具备承装资质"},
                    {"seq": 2, "item_label": "符合性审查", "description": "响应文件齐全"},
                ],
                "commerce_scores": [
                    {"title": "业绩", "criteria": "近三年同类业绩", "score_value": 10},
                ],
            },
        )
        rows2 = persist_commercial_draft(db, project, preserve_confirmed=True)
        db.commit()
        confirmed = next(r for r in rows2 if r.id == target.id)
        assert confirmed.status == STATUS_CONFIRMED
        assert "人工确认内容" in confirmed.content_markdown
        draft_rows = [r for r in rows2 if r.status == STATUS_DRAFT]
        assert draft_rows
    finally:
        db.close()


def test_commercial_status_includes_draft_count():
    init_db()
    db = SessionLocal()
    try:
        project = _seed_project(db)
        set_bid_scope(project, True)
        persist_commercial_draft(db, project)
        db.commit()
        from services.commercial_bid_service import commercial_status

        status = commercial_status(db, project)
        assert status["section_count"] > 0
        assert status["draft_count"] == status["section_count"]
        assert status["confirmed_count"] == 0
    finally:
        db.close()


def test_validate_commercial_export_ready_blocks_drafts():
    init_db()
    db = SessionLocal()
    try:
        project = _seed_project(db)
        set_bid_scope(project, True)
        rows = persist_commercial_draft(db, project)
        db.commit()
        from services.commercial_bid_service import validate_commercial_export_ready

        with pytest.raises(ValueError, match="未确认"):
            validate_commercial_export_ready(db, project, allow_draft=False)

        draft_count = validate_commercial_export_ready(db, project, allow_draft=True)
        assert draft_count == len(rows)

        update_commercial_section(
            db, project.id, rows[0].id, status=STATUS_CONFIRMED
        )
        db.commit()
        # 仍有其余草稿
        with pytest.raises(ValueError, match="未确认"):
            validate_commercial_export_ready(db, project, allow_draft=False)

        for row in rows[1:]:
            update_commercial_section(
                db, project.id, row.id, status=STATUS_CONFIRMED
            )
        db.commit()
        assert validate_commercial_export_ready(db, project, allow_draft=False) == 0
    finally:
        db.close()


def test_append_commercial_markdown_converts_table_to_word_table():
    doc = Document()
    md = (
        "## 资格对照\n\n"
        "说明如下：\n"
        "| 序号 | 审查项 | 响应 |\n"
        "|---|---|---|\n"
        "| 1 | 资质 | 完全响应 |\n"
        "| 2 | 业绩 | 详见附件 |\n"
        "以上为资格审查响应。"
    )
    _append_commercial_markdown_to_doc(doc, md)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.rows[0].cells[0].text == "序号"
    assert table.rows[1].cells[2].text == "完全响应"
    assert table.rows[2].cells[1].text == "业绩"
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "说明如下：" in texts
    assert "以上为资格审查响应。" in texts


def test_append_commercial_markdown_keeps_orphan_pipe_lines():
    """非标准 | 行不得静默丢弃，应保留为段落。"""
    doc = Document()
    _append_commercial_markdown_to_doc(doc, "| 这不是标准表格行\n正文继续")
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "| 这不是标准表格行" in texts
    assert "正文继续" in texts
    assert len(doc.tables) == 0
