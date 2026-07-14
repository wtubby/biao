"""自定义排版参数。"""

from docx import Document
from docx.oxml.ns import qn

from db.models import Project
from services.generation_config import update_generation_config
from services.typesetting_config import (
    DEFAULT_TYPESETTING,
    TypesettingNumbering,
    _register_custom_numbering,
    _resolve_lvl_text,
    apply_typesetting_styles,
    default_typesetting,
    get_typesetting,
    normalize_typesetting,
)


def test_default_typesetting_matches_shenjuan_style():
    ts = default_typesetting()
    assert ts["h1"]["font"] == "黑体"
    assert ts["h1"]["font_size"] == "小三"
    assert ts["h1"]["align"] == "center"
    assert ts["h1"]["number_format"] == "cn_dun"
    assert ts["h2"]["number_format"] == "cn_paren_dun"
    assert ts["body"]["first_line_indent"] == 2


def test_normalize_typesetting_clamps_indent():
    raw = {"body": {"first_line_indent": 99, "font": "未知字体"}}
    ts = normalize_typesetting(raw)
    assert ts["body"]["first_line_indent"] == 8
    assert ts["body"]["font"] == "宋体"


def test_get_typesetting_from_project():
    project = Project(id="p1", extra_params="{}")
    update_generation_config(project, typesetting={"h1": {"align": "right"}})
    ts = get_typesetting(project)
    assert ts["h1"]["align"] == "right"
    assert ts["h2"]["number_format"] == DEFAULT_TYPESETTING["h2"]["number_format"]


def test_resolve_lvl_text_single_level_uses_current_ilvl():
    assert _resolve_lvl_text("{n}、", 0) == "%1、"
    assert _resolve_lvl_text("（{n}）、", 1) == "（%2）、"
    assert _resolve_lvl_text("{n}、", 2) == "%3、"


def test_resolve_lvl_text_multi_level_shifts_with_ilvl():
    assert _resolve_lvl_text(2, 1) == "%1.%2"
    assert _resolve_lvl_text(2, 3) == "%3.%4"
    assert _resolve_lvl_text(3, 2) == "%1.%2.%3"
    assert _resolve_lvl_text(4, 4) == "%2.%3.%4.%5"


def test_register_custom_numbering_lvl_text_matches_ilvl():
    doc = Document()
    ts = default_typesetting()
    # 故意把 h1 设为原先写死 %2 的格式，验证现在绑定到 ilvl0
    ts["h1"]["number_format"] = "cn_paren_dun"
    ts["h3"]["number_format"] = "cn_dun"
    _register_custom_numbering(doc, ts)

    numbering_elm = doc.part.numbering_part.element
    abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
    abstract_num = abstract_nums[-1]

    lvl_texts: dict[int, str] = {}
    for lvl in abstract_num.findall(qn("w:lvl")):
        ilvl = int(lvl.get(qn("w:ilvl")))
        lvl_text_el = lvl.find(qn("w:lvlText"))
        lvl_texts[ilvl] = lvl_text_el.get(qn("w:val")) if lvl_text_el is not None else ""

    assert lvl_texts[0] == "（%1）、"
    assert lvl_texts[1] == "（%2）、"
    assert lvl_texts[2] == "%3、"
    assert lvl_texts[3] == "%3.%4"


def test_typesetting_numbering_registers_num_pr():
    doc = Document()
    ts = default_typesetting()
    numbering = TypesettingNumbering(doc, ts)
    assert numbering.enabled
    para = doc.add_heading("测试标题", level=1)
    numbering.apply(para, 1)
    p_pr = para._p.pPr
    assert p_pr is not None
    from docx.oxml.ns import qn
    assert p_pr.find(qn("w:numPr")) is not None


def test_apply_typesetting_styles_body_indent():
    doc = Document()
    para = doc.add_paragraph("正文段落示例。")
    ts = default_typesetting()
    apply_typesetting_styles(doc, ts)
    assert para.paragraph_format.first_line_indent is not None
    assert para.paragraph_format.first_line_indent.pt == 24  # 2 字符 × 小四 12pt
