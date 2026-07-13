"""章节标题多级编号。"""

import json
import zipfile
from types import SimpleNamespace

from docx import Document

from docx.oxml.ns import qn

from services.numbering_service import (
    HeadingNumbering,
    list_heading_numbering_presets,
    resolve_heading_numbering_preset,
)


def _paragraph_num_pr(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    ilvl_el = num_pr.find(qn("w:ilvl"))
    num_id_el = num_pr.find(qn("w:numId"))
    return {
        "ilvl": int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else None,
        "numId": int(num_id_el.get(qn("w:val"))) if num_id_el is not None else None,
    }


def test_heading_numbering_applies_num_pr_by_level():
    doc = Document()
    numbering = HeadingNumbering(doc, "decimal")
    assert numbering.enabled

    h1 = doc.add_heading("总体部署", level=1)
    h2 = doc.add_heading("施工安排", level=2)
    numbering.apply(h1, 1)
    numbering.apply(h2, 2)

    assert _paragraph_num_pr(h1) == {"ilvl": 0, "numId": numbering._num_id}
    assert _paragraph_num_pr(h2) == {"ilvl": 1, "numId": numbering._num_id}


def test_heading_numbering_none_skips_num_pr():
    doc = Document()
    numbering = HeadingNumbering(doc, "none")
    assert not numbering.enabled

    heading = doc.add_heading("无编号标题", level=1)
    numbering.apply(heading, 1)
    assert _paragraph_num_pr(heading) is None


def test_register_decimal_multilevel_in_numbering_xml():
    doc = Document()
    numbering = HeadingNumbering(doc, "decimal")
    doc.add_heading("章", level=1)
    numbering.apply(doc.paragraphs[-1], 1)

    buf = __import__("io").BytesIO()
    doc.save(buf)
    with zipfile.ZipFile(buf) as zf:
        numbering_xml = zf.read("word/numbering.xml").decode("utf-8")

    assert "%1.%2.%3" in numbering_xml
    assert f'w:numId="{numbering._num_id}"' in numbering_xml


def test_resolve_heading_numbering_preset_project_override():
    project = SimpleNamespace(
        extra_params=json.dumps({"heading_numbering_preset": "chapter_cn"})
    )
    assert resolve_heading_numbering_preset(project) == "chapter_cn"


def test_resolve_heading_numbering_preset_invalid_falls_back_decimal():
    project = SimpleNamespace(extra_params=json.dumps({"heading_numbering_preset": "unknown"}))
    assert resolve_heading_numbering_preset(project) == "decimal"


def test_heading_numbering_uses_emit_level_when_container_skipped():
    from types import SimpleNamespace

    from docx.oxml.ns import qn

    from services.assembler_service import assemble_document

    project = SimpleNamespace(
        id="p1",
        name="测试",
        voltage_level="",
        capacity="",
        location="",
        duration_days=90,
        extra_params=None,
    )
    chapters = [
        SimpleNamespace(
            id="1",
            parent_id=None,
            title="施工组织",
            level=1,
            is_leaf=0,
            generated_content=None,
            review_status="green",
            sort_order=1,
        ),
        SimpleNamespace(
            id="1.1",
            parent_id="1",
            title="土建工程",
            level=2,
            is_leaf=0,
            generated_content=None,
            review_status="green",
            sort_order=2,
        ),
        SimpleNamespace(
            id="1.1.1",
            parent_id="1.1",
            title="基础施工",
            level=3,
            is_leaf=1,
            generated_content="正文",
            review_status="green",
            sort_order=3,
        ),
    ]

    with __import__("tempfile").TemporaryDirectory() as tmp:
        import services.assembler_service as asm

        old_out = asm.OUTPUT_DIR
        asm.OUTPUT_DIR = tmp
        try:
            path = assemble_document(project, chapters)
            doc = Document(str(path))
        finally:
            asm.OUTPUT_DIR = old_out

    heading_paras = [
        p for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading")
        and p._p.pPr is not None
        and p._p.pPr.find(qn("w:numPr")) is not None
    ]
    assert len(heading_paras) == 2
    ilvls = []
    for p in heading_paras:
        num_pr = p._p.pPr.find(qn("w:numPr")) if p._p.pPr is not None else None
        ilvl_el = num_pr.find(qn("w:ilvl")) if num_pr is not None else None
        ilvls.append(int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else None)
    assert ilvls == [0, 1]


def test_list_heading_numbering_presets_includes_none_and_decimal():
    presets = list_heading_numbering_presets()
    ids = [p["id"] for p in presets]
    assert "none" in ids
    assert "decimal" in ids
    assert "chapter_cn" in ids

