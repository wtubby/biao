"""Word 装配服务单元测试，重点覆盖 Markdown 表格解析。"""

from docx import Document

import re as _re

from services.assembler_service import (
    _compute_heading_emit_levels,
    _detect_list_style,
    _insert_chart,
    _insert_markdown_list,
    _insert_markdown_table,
    _is_markdown_table,
    _next_caption_number,
    _parse_table_cells,
    _should_emit_heading,
    _write_text_block,
)


def test_parse_table_cells_strips_outer_pipes():
    assert _parse_table_cells("| a | b | c |") == ["a", "b", "c"]


def test_parse_table_cells_handles_escaped_pipe():
    assert _parse_table_cells(r"| a\|b | c |") == ["a|b", "c"]


def test_is_markdown_table_detects_valid_table():
    lines = ["| 设备 | 型号 | 数量 |", "|---|---|---|", "| 主变 | SZ11 | 2 |"]
    assert _is_markdown_table(lines)


def test_is_markdown_table_rejects_plain_text():
    lines = ["这是一段普通正文，包含 | 字符但不是表格。", "第二行普通文字。"]
    assert not _is_markdown_table(lines)


def test_insert_markdown_table_builds_real_table():
    doc = Document()
    lines = [
        "| 设备名称 | 型号 | 数量 |",
        "|---|---|---|",
        "| 主变压器 | SZ11-50000/110 | 2 |",
        "| GIS | ZF12 | 12 |",
    ]
    consumed = _insert_markdown_table(doc, lines)
    assert consumed == 4
    tables = doc.tables
    assert len(tables) == 1
    table = tables[0]
    assert len(table.rows) == 3  # 表头 + 2 行数据
    assert len(table.columns) == 3
    assert table.rows[0].cells[0].text == "设备名称"
    assert table.rows[1].cells[0].text == "主变压器"
    assert table.rows[2].cells[2].text == "12"


def test_write_text_block_converts_table_and_keeps_trailing_text():
    doc = Document()
    content = (
        "本章设备清单如下：\n\n"
        "| 设备 | 数量 |\n"
        "|---|---|\n"
        "| 主变 | 2 |\n"
        "以上为主要设备统计。"
    )
    _write_text_block(doc, content, chapter_level=1)
    assert len(doc.tables) == 1
    paragraph_texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "本章设备清单如下：" in paragraph_texts
    assert "以上为主要设备统计。" in paragraph_texts


def test_should_emit_heading_rules():
    from types import SimpleNamespace

    leaf = SimpleNamespace(is_leaf=1, level=3, generated_content="正文")
    l1_container = SimpleNamespace(is_leaf=0, level=1, generated_content=None)
    l2_empty = SimpleNamespace(is_leaf=0, level=2, generated_content=None)
    with_content = SimpleNamespace(is_leaf=0, level=2, generated_content="概述")

    assert _should_emit_heading(leaf) is True
    assert _should_emit_heading(l1_container) is True
    assert _should_emit_heading(l2_empty) is False
    assert _should_emit_heading(with_content) is True


def test_compute_heading_emit_levels_skips_empty_containers():
    from types import SimpleNamespace

    chapters = [
        SimpleNamespace(id="1", parent_id=None, level=1, is_leaf=0, generated_content=None),
        SimpleNamespace(id="1.1", parent_id="1", level=2, is_leaf=0, generated_content=None),
        SimpleNamespace(id="1.1.1", parent_id="1.1", level=3, is_leaf=1, generated_content="A"),
        SimpleNamespace(id="1.1.2", parent_id="1.1", level=3, is_leaf=1, generated_content="B"),
    ]
    levels = _compute_heading_emit_levels(chapters)
    assert levels == {"1": 1, "1.1.1": 2, "1.1.2": 2}


def test_detect_list_style_bullet_and_number():
    assert _detect_list_style("- 第一项") == "List Bullet"
    assert _detect_list_style("1. 第一项") == "List Number"
    assert _detect_list_style("普通文字") is None


def test_insert_markdown_list_consumes_consecutive_items_only():
    doc = Document()
    lines = ["- 基础开挖", "- 设备就位", "- 二次接线", "以上为主要工序。"]
    consumed = _insert_markdown_list(doc, lines, "List Bullet")
    assert consumed == 3
    paragraphs = [p.text for p in doc.paragraphs]
    assert "基础开挖" in paragraphs
    assert "以上为主要工序。" not in paragraphs  # 未被列表吞掉，留给外层继续处理


def test_write_text_block_handles_list_then_trailing_text():
    doc = Document()
    content = "施工工序如下：\n\n1. 基础开挖\n2. 设备就位\n以上为主要工序。"
    _write_text_block(doc, content, chapter_level=1)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "基础开挖" in texts
    assert "以上为主要工序。" in texts


def test_write_text_block_handles_intro_then_list_single_newline():
    doc = Document()
    content = "施工工序如下：\n1. 基础开挖\n2. 设备就位\n以上为主要工序。"
    _write_text_block(doc, content, chapter_level=1)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "施工工序如下：" in texts
    assert "基础开挖" in texts
    assert "以上为主要工序。" in texts


def test_next_caption_number_increments_per_kind_independently():
    counters: dict = {}
    assert _next_caption_number(counters, "figure") == 1
    assert _next_caption_number(counters, "figure") == 2
    assert _next_caption_number(counters, "table") == 1


def test_insert_chart_gantt_block_format_object():
    """块格式 [GANTT_DATA]\\n{tasks:...} 应渲染为图片而非原文。"""
    from chart.chart_service import CHART_PATTERN
    from services.assembler_service import _write_content_in_order

    content = """进度计划如下：

```json
[GANTT_DATA]
{
  "tasks": [
    {"name": "基础施工", "start": 0, "end": 30},
    {"name": "设备安装", "start": 30, "end": 90}
  ]
}
```

后续说明文字。"""
    doc = Document()
    _write_content_in_order(doc, content, 2, [], 180, {})
    assert not any("GANTT_DATA" in p.text for p in doc.paragraphs)
    assert len(doc.inline_shapes) == 1
    captions = [p.text for p in doc.paragraphs if p.text.startswith("图")]
    assert captions == ["图1 施工进度横道图"]


def test_insert_chart_gantt_now_inserts_picture_not_table():
    from chart.chart_service import CHART_PATTERN

    doc = Document()
    content = '[GANTT_DATA: [{"工序": "基础施工", "开始第几天": 1, "持续天数": 5}]]'
    match = next(CHART_PATTERN.finditer(content))
    _insert_chart(doc, match, temp_files=[], duration=30, counters={})
    assert len(doc.tables) == 0
    assert len(doc.inline_shapes) == 1
    captions = [p.text for p in doc.paragraphs if p.text.startswith("图")]
    assert captions == ["图1 施工进度横道图"]


def test_insert_chart_smart_data_skips_empty_array():
    """SMART_DATA 空数组不插表、不插图注，且不消耗表编号。"""
    from chart.chart_service import CHART_PATTERN

    doc = Document()
    counters: dict = {}
    empty = next(CHART_PATTERN.finditer("[SMART_DATA: []]"))
    _insert_chart(doc, empty, temp_files=[], duration=30, counters=counters)
    assert len(doc.tables) == 0
    assert not any(p.text.startswith("表") for p in doc.paragraphs)
    assert counters.get("table", 0) == 0

    filled = next(
        CHART_PATTERN.finditer('[SMART_DATA: [{"title": "工期", "desc": "180日历天"}]]')
    )
    _insert_chart(doc, filled, temp_files=[], duration=30, counters=counters)
    assert len(doc.tables) == 1
    captions = [p.text for p in doc.paragraphs if p.text.startswith("表")]
    assert captions == ["表1 要点对照表"]
